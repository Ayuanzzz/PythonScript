import docx
import os
import comtypes.client
import aspose.words as aw


def get_pictures(word_path, result_path):
    try:
        doc = docx.Document(word_path)
        first_page_images = []
        page_break_found = False

        for rel in doc.part.rels:
            rel = doc.part.rels[rel]
            if page_break_found:
                break  # 停止在页面分隔符后继续读取

            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                first_page_images.append(image_data)

            for paragraph in doc.paragraphs:
                if '\x0C' in paragraph.text:  # 检测页面分隔符
                    page_break_found = True
                    break

        for i, image_data in enumerate(first_page_images):
            img_name = os.path.basename(word_path).replace("_.docx", ".jpg")
            with open(os.path.join(result_path, img_name), "wb") as f:
                f.write(image_data)
                print(f"提取图片: {img_name}")

    except Exception as e:
        print(f"发生错误：{e}")


def docToDocx(doc_file, docx_file):
    # 启动 Microsoft Word 应用程序
    word = comtypes.client.CreateObject("Word.Application")
    # 打开 .doc 文件
    doc = word.Documents.Open(doc_file)
    # 保存为 .docx 文件
    doc.SaveAs2(docx_file, FileFormat=16)  # 16 表示 .docx 格式
    # 关闭 .doc 文件和 Microsoft Word 应用程序
    doc.Close()
    word.Quit()


if __name__ == '__main__':
    # result_path = r"D:\Python\test\doc"
    result_path = input("请输入word文件夹路径: ")
    for root, dirs, files in os.walk(result_path):
        for file in files:
            if file.endswith(".doc"):
                # 这里执行你的操作，例如打印文件路径
                doc_file_path = os.path.join(root, file)
                print(doc_file_path)
                try:
                    docx_path_old = doc_file_path.replace(".doc", ".docx")
                    docToDocx(doc_file_path, docx_path_old)
                    docx_path = docx_path_old.replace(".docx", "_.docx")
                    doc = aw.Document(docx_path_old)
                    extractedPage = doc.extract_pages(0, 1)
                    extractedPage.save(docx_path)
                    get_pictures(docx_path, result_path)
                    os.remove(docx_path)
                    os.remove(docx_path_old)
                except Exception as e:
                    print(f"发生错误：{e}")

print("\n提取完成")
input("\n按任意键退出")
