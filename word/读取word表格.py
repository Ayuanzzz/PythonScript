import re
import cn2an
import openpyxl
import os

from docx import Document


def read_specific_cell(doc_path, page, rowIndex, cellIndex):
    try:
        doc = Document(doc_path)

        # 遍历文档中的所有表格
        for table_index, table in enumerate(doc.tables):
            if table_index == page:  # 仅处理第一张表
                for row_index, row in enumerate(table.rows):
                    if row_index == rowIndex:
                        for cell_index, cell in enumerate(row.cells):
                            if cell_index == cellIndex:
                                return cell.text

    except Exception as e:
        print(f"Error: {e}")


def hasOrNone(has_status):
    if has_status == "□有☑无":
        return 0
    elif has_status == "☑有□无":
        return 1


def yesOrNo(yn_status):
    if yn_status == "□是☑否":
        return 0
    elif yn_status == "☑是□否":
        return 1


def readTabel(doc_path, page):
    tmp_str = read_specific_cell(doc_path, page, 0, 0)
    # 权利人
    # print("\nPerson")
    person = ""
    pattern = re.compile(r'权利人：(.*)')
    match = pattern.search(tmp_str)
    if match:
        person = match.group(1)
    else:
        print("权利人")
    # print(person)
    # 组
    # print("\nNum")
    num = 0
    # num_pattern = re.compile(r'村(.*?)组'
    num_pattern = re.compile(r'社区(.*?)组')
    num_match = num_pattern.search(tmp_str)
    if num_match:
        try:
            num = cn2an.cn2an(num_match.group(1).strip())
        except:
            num = int(num_match.group(1).strip())
    else:
        print("未找到组别")
    # print(num)
    # 地籍要素有无错漏
    # print("\nDJYS")
    tmp_DJYS = read_specific_cell(doc_path, page, 2, 1)
    # print(tmp_DJYS)
    DJYS = hasOrNone(tmp_DJYS)
    # print(DJYS)
    # 地籍图上表示的地籍要素与地籍调查结果是否一致
    # print("\nDJT")
    tmp_DJT = read_specific_cell(doc_path, page, 3, 1)
    # print(tmp_DJT)
    DJT = yesOrNo(tmp_DJT)
    # print(DJT)
    # 房屋结构有无错漏
    # print("\nFWJG")
    tmp_FWJG = read_specific_cell(doc_path, page, 4, 1)
    # print(tmp_FWJG)
    FWJG = hasOrNone(tmp_FWJG)
    # print(FWJG)
    # 房屋层数有无错漏
    # print("\nFWCS")
    tmp_FWCS = read_specific_cell(doc_path, page, 5, 1)
    # print(tmp_FWCS)
    FWCS = hasOrNone(tmp_FWCS)
    # print(FWCS)
    # 权利人和户籍信息是否正确
    # print("\nQLR")
    tmp_QLR = read_specific_cell(doc_path, page, 6, 1)
    # print(tmp_QLR)
    QLR = yesOrNo(tmp_QLR)
    # print(QLR)
    # 本宗地全貌调查是否开展
    # print("\nBZD")
    tmp_BZD = read_specific_cell(doc_path, page, 7, 1)
    # print(tmp_BZD)
    BZD = yesOrNo(tmp_BZD)
    # print(BZD)
    # 其他问题描述
    # print("\nQTWT")
    tmp_QTWT = read_specific_cell(doc_path, page, 8, 1)
    # print(tmp_QTWT)
    QTWT = hasOrNone(tmp_QTWT)
    # print(QTWT)
    # 备注
    remake_list = [read_specific_cell(doc_path, page, 2, 2), read_specific_cell(doc_path, page, 3, 2),
                   read_specific_cell(
                       doc_path, page, 4, 2), read_specific_cell(doc_path, page, 5, 2),
                   read_specific_cell(doc_path, page, 6, 2), read_specific_cell(doc_path, page, 7, 2),
                   read_specific_cell(doc_path, page, 8, 2)]
    remake = ",".join(element for element in remake_list if element)

    table_list = [num, person, DJYS, DJT, FWJG, FWCS, QLR, BZD, QTWT, remake, country_name, village_name]
    # print(table_list)
    return table_list


def calcNum(input_num):
    if input_num == 0:
        return ["", "1"]
    if input_num == 1:
        return ["1", ""]
    else:
        return ["", ""]


if __name__ == "__main__":
    # 写入excel
    excel_path = r"D:\Python\test\wordToExcel\excel\target.xlsx"
    new_path = r"D:\Python\test\wordToExcel\excel\新棉街道.xlsx"
    workbook = openpyxl.load_workbook(excel_path)
    sheet = workbook.active

    folder_path = r"D:\Python\test\记录表装订\记录表装订\1-检查报告及记录\新棉街道-362"
    count = 4
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith(".docx"):
                # 遍历文档中的所有表格
                doc_path = os.path.join(root, file)
                print("++++++++++++++++++++++++++++++++++++++++++++++++++++")
                print(doc_path)
                doc = Document(doc_path)
                # 村名
                path, file_name = os.path.split(doc_path)
                village_name = file_name.split("-")[0]
                # 乡镇名
                parent_path, folder_name = os.path.split(path)
                country_name = folder_name.split("-")[0]

                for table_index, table in enumerate(doc.tables):
                    # print(table_index)
                    info_list = readTabel(doc_path, table_index)

                    # 写入乡镇名
                    sheet.cell(row=count, column=2, value=info_list[10])
                    # 写入村民
                    sheet.cell(row=count, column=3, value=info_list[11])
                    # 写入组别
                    sheet.cell(row=count, column=4, value=info_list[0])
                    # 写入权利人
                    sheet.cell(row=count, column=5, value=info_list[1])
                    # 写入地籍要素有无错漏
                    sheet.cell(row=count, column=6, value=calcNum(info_list[2])[0])
                    sheet.cell(row=count, column=7, value=calcNum(info_list[2])[1])
                    # 地籍图上表示的地籍要素与地籍调查结果是否一致
                    sheet.cell(row=count, column=8, value=calcNum(info_list[3])[0])
                    sheet.cell(row=count, column=9, value=calcNum(info_list[3])[1])
                    # 房屋结构有无错漏
                    sheet.cell(row=count, column=10, value=calcNum(info_list[4])[0])
                    sheet.cell(row=count, column=11, value=calcNum(info_list[4])[1])
                    # 房屋层数有无错漏
                    sheet.cell(row=count, column=12, value=calcNum(info_list[5])[0])
                    sheet.cell(row=count, column=13, value=calcNum(info_list[5])[1])
                    # 权利人和户籍信息是否正确
                    sheet.cell(row=count, column=14, value=calcNum(info_list[6])[0])
                    sheet.cell(row=count, column=15, value=calcNum(info_list[6])[1])
                    # 本宗地全貌调查是否开展
                    sheet.cell(row=count, column=16, value=calcNum(info_list[7])[0])
                    sheet.cell(row=count, column=17, value=calcNum(info_list[7])[1])
                    # 其他问题描述
                    sheet.cell(row=count, column=18, value=calcNum(info_list[8])[0])
                    sheet.cell(row=count, column=19, value=calcNum(info_list[8])[1])
                    # 写入备注
                    sheet.cell(row=count, column=20, value=info_list[9])
                    # print("---------")
                    count += 1

                # 保存文件
                workbook.save(new_path)
