import pandas as pd
import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

directory = input("输入文件夹路径:")
zj_time = input("\n输入指界时间(年-月-日):")
hz_time = input("\n输入回执时间(年-月-日):")
person = input("\n输入经办人名称:")
print("--------------------------------")
print("\n程序执行中...\n")


def preserve_formatting(original_run, new_run):
    new_run.bold = original_run.bold
    new_run.italic = original_run.italic
    new_run.font.size = original_run.font.size
    new_run.font.color.rgb = original_run.font.color.rgb


def modify_docx(template_path, output_path, header_text, footer_text, location, zj_time, hz_time):
    doc = Document(template_path)

    # Insert header
    header = doc.sections[0].header
    header.paragraphs[0].text = header_text
    header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Insert footer
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = footer_text
    footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    zj_parts = zj_time.split("-")
    hz_parts = hz_time.split("-")

    for paragraph in doc.paragraphs:
        for to_replace, replacement_text in [("土地坐落：", "土地坐落：" + location),
                                             ("集合地点：", "集合地点：" + location),
                                             ("指界时间：",
                                              f"指界时间：{zj_parts[0]}  年   {zj_parts[1]} 月   {zj_parts[2]} 日      午      时"),
                                             ("通知日期：",
                                              f"通知日期：{zj_parts[0]}  年   {zj_parts[1]} 月   {zj_parts[2]} 日"),
                                             ("经办人：", f"经办人：{person}")]:
            if to_replace in paragraph.text:
                # run = paragraph.runs[0]
                for run in paragraph.runs:
                    run.text = run.text.replace(to_replace, replacement_text)
                    # Preserve original formatting
                    new_run = paragraph.add_run()
                    preserve_formatting(run, new_run)

    tables = doc.tables
    last_table = tables[-1]
    last_row = last_table.rows[-1]
    for cell in last_row.cells:
        if "ffff年ff月ff日" in cell.text:
            cell.text = cell.text.replace("ffff年ff月ff日", f"{hz_parts[0]}年{hz_parts[1]}月{hz_parts[2]}日")

    # Save the modified document
    doc.save(output_path)


def read_xls(xls_path):
    # Read the Excel file
    df = pd.read_excel(xls_path)
    return df


for file in os.listdir(directory):
    if file.endswith(".xls"):
        xlsname = file.replace(".xls", "")
        print(xlsname)
        xls_path = os.path.join(directory, file)
        output_path = os.path.join(directory, xlsname)
        os.makedirs(output_path)

        data_frame = read_xls(xls_path)

        all_sheets = pd.read_excel(xls_path, sheet_name=None)
        for sheet_name, sheet_data in all_sheets.items():
            print(f"Sheet: {sheet_name}")

            for index, row in sheet_data.iterrows():
                num = row["序号"]
                doc_name = row["不动产权证号"]
                people_name = row["姓名"]
                location = row["现坐落（地址）"]

                try:
                    template_path = "template.docx"
                    filename = doc_name + "-" + people_name + ".docx"
                    header_text = people_name
                    footer_text = header_text
                    output_xls_path = os.path.join(output_path, filename)
                    modify_docx(template_path, output_xls_path, header_text, footer_text, location, zj_time, hz_time)
                except:
                    print(num)

print("\n处理完成")
input("\n按任意键退出")
